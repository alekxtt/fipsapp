import datetime
import locale
import requests
import os
from bs4 import BeautifulSoup as bs
from docx import Document as Word_Document
from docx.shared import Inches as Word_Inches

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.mail import send_mail, EmailMessage
from django.conf import settings
from django.shortcuts import redirect, render

from .models import IpItem
from .forms import IpItemForm
from .text import *


@login_required(login_url='profile:login')
def items_list(request):
    pass


@login_required(login_url='profile:login')
def item_add(request):
    template = 'ipitem/ipitem_form.html'
    form = IpItemForm(request.POST or None)
    type_request = 'new'
    if form.is_valid():
        ipitem = form.save(commit=False)
        ipitem.title = ipitem.title.upper()
        ipitem.owner = request.user.profile
        ipitem.save()
        messages.success(request, 'Объект был успешно добавлен')
        return redirect('profile:profile')
    else:
        messages.error(request, 'Проверьте введённые данные')
    context = {
        'form': form,
        'type_request': type_request,
    }
    return render(request, template, context)


@login_required(login_url='profile:login')
def item_edit(request, pk):
    ipitem = IpItem.objects.get(id=pk)
    type_request = 'edit'
    if request.user.profile != ipitem.owner:
        return redirect('profile:profile')
    template = 'ipitem/ipitem_form.html'
    form = IpItemForm(instance=ipitem)
    print(form)
    if request.method == "POST":
        form = IpItemForm(request.POST, instance=ipitem)
        if form.is_valid():
            ipitem = form.save(commit=False)
            ipitem.title = ipitem.title.upper()
            ipitem.save()
            messages.success(request, 'Объект был успешно изменён')
            return redirect('profile:profile')
    context = {
        'type_request': type_request,
        'form': form,
        'ipitem': ipitem,
    }
    return render(request, template, context)


@login_required(login_url='profile:login')
def item_info(request, pk):
    template = 'ipitem/ipitem_page.html'
    item = IpItem.objects.get(id=pk)
    type_of_item = ITEMS_FOR_CHOICE[item.type_of_item]
    context = {
        'item': item,
        'type_of_item': type_of_item,
    }
    return render(request, template, context)


@login_required(login_url='profile:login')
def item_delete_confirmation(request, pk):
    template = 'ipitem/ipitem_page_delete.html'
    item = IpItem.objects.get(id=pk)
    type_of_item = ITEMS_FOR_CHOICE[item.type_of_item]
    context = {
        'item': item,
        'type_of_item': type_of_item,
    }
    return render(request, template, context)


@login_required(login_url='profile:login')
def item_delete(request, pk):
    ipitem = IpItem.objects.get(id=pk)
    if request.user.profile != ipitem.owner:
        return redirect('profile:profile')
    ipitem.delete()
    return redirect('profile:profile')


@login_required(login_url='profile:login')
def item_check_on_fips(request, pk):
    ipitem = IpItem.objects.get(id=pk)
    template = (f'https://fips.ru/registers-doc-view/fips_servlet?DB'
                f'={ TYPE_OF_IP[ipitem.type_of_item] }&DocNumber='
                f'{ ipitem.number }&TypeFile=html')
    return redirect(template)


def create_attached_files_to_email(ipitem, period,
                                   date_for_sign, continue_of_letter):
    # created by https://python-docx.readthedocs.io/
    yearly_payment = TYPE_IT_PAYMENT[ipitem.type_of_item][str(period)]
    paymet_document = Word_Document()
    paymet_document.add_paragraph(yearly_payment)
    paymet_document_title = 'payment_'+str(ipitem.number)+'_.docx'
    paymet_document.save(paymet_document_title)
    letter_for_fips = Word_Document()
    paragraph_letter_for_fips = letter_for_fips.add_paragraph(TITLE)
    paragraph_letter_for_fips_format = paragraph_letter_for_fips.paragraph_format
    paragraph_letter_for_fips_format.left_indent = Word_Inches(3.5)
    paragraph_letter_for_fips_2 = letter_for_fips.add_paragraph(ADDRESS)
    paragraph_letter_for_fips_format_2 = paragraph_letter_for_fips_2.paragraph_format
    paragraph_letter_for_fips_format_2.left_indent = Word_Inches(3.5)
    letter_for_fips.add_paragraph(BEGGINING_OF_LETTER_PATENT
    +str(period)+continue_of_letter+str(ipitem.number))
    letter_for_fips.add_paragraph(SIGN+ipitem.signatory)
    paragraph_letter_for_fips_3 = letter_for_fips.add_paragraph(date_for_sign)
    paragraph_letter_for_fips_format_3 = paragraph_letter_for_fips_3.paragraph_format
    paragraph_letter_for_fips_format_3.left_indent = Word_Inches(0.5)
    letter_for_fips_title = 'letter_'+str(ipitem.number)+'_.docx'
    letter_for_fips.save(letter_for_fips_title)
    return paymet_document_title, letter_for_fips_title

def send_email(ipitem, pk, period, date_for_sign, continue_of_letter):
    # отправка электронного письма пользователю с инфой для продления
    subject = (f'Документы для продления: '
               f'{ ITEMS_FOR_CHOICE[ipitem.type_of_item] } { ipitem.number }')
    message = (f'Здравствуйте! \n Согласно вашего запроса направляем вам'
               f' документы для продления объекта: '
               f' { ITEMS_FOR_CHOICE[ipitem.type_of_item] }'
               f' { ipitem.number }.\n'
               f'Документы для продления доступны по ссылкам: \n'
               f'1) Письмо: http://127.0.0.1:8000/ipitem/item/'
               f'item_prepare_letter/{ pk }/no_send_email/ \n'
               f'2) Реквизиты для оплаты пошлины: https://'
               f'rospatent.gov.ru/content/uploadfiles/rekviziti-pril1.pdf \n'
               f'3) Назначения и суммы платежей: '
               f'http://127.0.0.1:8000/ipitem/item/item_yearly_payment/{ pk }/')
    # send_mail(
    #     subject,
    #     message,
    #     settings.EMAIL_HOST_USER,
    #     [ipitem.owner.email],
    #     fail_silently=False,
    # )
    email = EmailMessage(
                    subject,
                    message,
                    settings.EMAIL_HOST_USER,
                    [ipitem.owner.email],
                    )
    (payment_file_for_email,
     letter_file_for_email) = create_attached_files_to_email(ipitem,
     period, date_for_sign, continue_of_letter)
    email.attach_file(payment_file_for_email)
    email.attach_file(letter_file_for_email)
    email.send()
    if os.path.exists(payment_file_for_email) and os.path.exists(letter_file_for_email):
        os.remove(payment_file_for_email)
        os.remove(letter_file_for_email)
    else:
        print("The files do not exist")


def year_of_continue(pk):
    # рассчёт года продления патента
    ipitem = IpItem.objects.get(id=pk)
    continue_of_letter = CONTINUE_OF_LETTER_PATENT[ipitem.type_of_item]
    # язык даты в подписи письма на русском:
    locale.setlocale(locale.LC_ALL, 'ru_RU.utf8')
    # год, месяц, день даты начала действия патента - init
    ip_init_year = int(ipitem.initial_date.strftime("%Y"))
    month_init_year = int(ipitem.initial_date.strftime("%m"))
    day_init_year = int(ipitem.initial_date.strftime("%d"))
    # год, месяц, день текущего года
    this_year = int(datetime.date.today().strftime("%Y"))
    month_this_year = int(datetime.date.today().strftime("%m"))
    day_this_year = int(datetime.date.today().strftime("%d"))
    # количество дней с 1 января до дня в дате начала действия
    delta_init_year = ((datetime.date(ip_init_year,
                                      month_init_year,
                                      day_init_year)
                       - datetime.date(ip_init_year, 1, 1)))
    # количество дней с 1 января до дня текущего в текущем году
    delta_this_year = ((datetime.date(this_year,
                                      month_this_year,
                                      day_this_year)
                       - datetime.date(this_year, 1, 1)))
    # год на основе разниц текущего года и init года
    period = this_year - ip_init_year
    # прибавка года на основе дат, если дата до init день, то +1, иначе +2
    if delta_init_year > delta_this_year:
        period += 1
    else:
        period += 2
    # дата подписания на письмо
    FORMAT = "%d.%m.%Y"
    date_for_sign = datetime.date.today().strftime(FORMAT)
    return period, date_for_sign, continue_of_letter, ipitem


@login_required(login_url='profile:login')
def item_prepare_letter_no_send_email(request, pk):
    period, date_for_sign, continue_of_letter, ipitem = year_of_continue(pk)
    context = {
        'ipitem': ipitem,
        'title': TITLE,
        'address': ADDRESS,
        'beginnig': BEGGINING_OF_LETTER_PATENT,
        'continue': continue_of_letter,
        'sign': SIGN,
        'date_for_sign': date_for_sign,
        'period': period,
    }
    template = 'ipitem/ipitem_letter.html'
    return render(request, template, context)


@login_required(login_url='profile:login')
def item_prepare_letter(request, pk):
    period, date_for_sign, continue_of_letter, ipitem = year_of_continue(pk)
    context = {
        'ipitem': ipitem,
        'title': TITLE,
        'address': ADDRESS,
        'beginnig': BEGGINING_OF_LETTER_PATENT,
        'continue': continue_of_letter,
        'sign': SIGN,
        'date_for_sign': date_for_sign,
        'period': period,
    }
    template = 'ipitem/ipitem_letter.html'
    send_email(ipitem, pk, period, date_for_sign, continue_of_letter)
    return render(request, template, context)


@login_required(login_url='profile:login')
def item_payments(request, pk):
    period, date_for_sign, continue_of_letter, ipitem = year_of_continue(pk)
    yearly_payment = TYPE_IT_PAYMENT[ipitem.type_of_item][str(period)]
    context = {
        'yearly_payment': yearly_payment,
    }
    template = 'ipitem/ipitem_payments.html'
    return render(request, template, context)


def parsing_fips(ipitem):
    # создано с помощью import requests
    template_for_pars = (f'https://fips.ru/registers-doc-view/fips_servlet?DB'
                         f'={ TYPE_OF_IP[ipitem.type_of_item] }&DocNumber='
                         f'{ ipitem.number }&TypeFile=html')
    base_pars = bs(requests.get(template_for_pars).text).findAll('p')
    initial_date_pars = base_pars[4].find('b').text
    FORMAT = "%d.%m.%Y"
    dt = datetime.datetime.strptime(initial_date_pars, FORMAT)
    initial_date_pars = str(dt.year)+'-'+str(dt.month)+'-'+str(dt.day)
    title_pars = base_pars[10].find('b').text
    description_pars = base_pars[12].text[:300]
    return initial_date_pars, title_pars, description_pars


@login_required(login_url='profile:login')
def fips_pars(request):
    type_request = 'pars'
    form = IpItemForm(request.POST or None)
    if form.is_valid():
        ipitem = form.save(commit=False)
        initial_date_pars, title_pars, description_pars = parsing_fips(ipitem)
        ipitem.initial_date = initial_date_pars
        ipitem.title = title_pars
        ipitem.title = ipitem.title.upper()
        ipitem.description = description_pars
        ipitem.owner = request.user.profile
        ipitem.save()
        messages.success(request, 'Объект был успешно добавлен')
        return redirect('profile:profile')
    else:
        messages.error(request, 'Проверьте введённые данные')
    context = {
        'type_request': type_request,
        'form': form,
    }
    template = 'ipitem/ipitem_form.html'
    return render(request, template, context)
